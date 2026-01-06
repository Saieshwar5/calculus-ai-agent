"""
RAG (Retrieval-Augmented Generation) Service.
Handles document processing, storage in Qdrant, and retrieval for context-aware responses.
"""
import os
import io
import uuid
import logging
import traceback
from typing import List, Dict, Any, Optional, Tuple
from datetime import datetime
from dataclasses import dataclass

from qdrant_client import QdrantClient
from qdrant_client.http import models
from qdrant_client.http.models import (
    Distance,
    VectorParams,
    PointStruct,
    Filter,
    FieldCondition,
    MatchValue
)
from fastapi import UploadFile
import chardet

from app.long_term_memory.shared.embedding import get_embedding_service

# Configure logging to show in console
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Qdrant configuration
QDRANT_URL = os.getenv(
    "QUADRANT_ENDPOINT",
    ""
)
QDRANT_API_KEY = os.getenv(
    "QUADRANT_API",
    ""
)

# Embedding dimension for text-embedding-3-small
EMBEDDING_DIMENSION = 1536

# Chunking configuration
CHUNK_SIZE = 2000  # characters (approximately 500 tokens)
CHUNK_OVERLAP = 400  # characters (approximately 100 tokens)


@dataclass
class DocumentChunk:
    """Represents a chunk of a document with metadata."""
    text: str
    filename: str
    file_type: str
    chunk_index: int
    total_chunks: int
    created_at: datetime


@dataclass
class RetrievedChunk:
    """Represents a retrieved chunk from Qdrant with score."""
    text: str
    filename: str
    score: float
    chunk_index: int


class QdrantClientManager:
    """
    Singleton manager for Qdrant client connection.
    Handles collection creation and management per user.
    """
    _instance: Optional['QdrantClientManager'] = None
    _client: Optional[QdrantClient] = None
    
    def __new__(cls):
        if cls._instance is None:
            cls._instance = super().__new__(cls)
        return cls._instance
    
    def __init__(self):
        if self._client is None:
            print("[RAG] Initializing Qdrant client...")
            print(f"[RAG] URL: {QDRANT_URL}")
            try:
                self._client = QdrantClient(
                    url=QDRANT_URL,
                    api_key=QDRANT_API_KEY,
                    timeout=30.0,  # Reduced timeout
                    prefer_grpc=False  # Use REST API instead of gRPC
                )
                # Test connection immediately
                print("[RAG] Testing Qdrant connection...")
                collections = self._client.get_collections()
                print(f"[RAG] ✅ Qdrant client initialized. Found {len(collections.collections)} collections")
            except Exception as e:
                print(f"[RAG] ❌ Failed to initialize Qdrant client: {e}")
                traceback.print_exc()
                raise
    
    @property
    def client(self) -> QdrantClient:
        """Get the Qdrant client instance."""
        return self._client
    
    def get_collection_name(self, user_id: str) -> str:
        """Generate collection name for a user."""
        # Sanitize user_id for collection name (replace special chars)
        safe_user_id = user_id.replace("-", "_").replace("@", "_at_").replace(".", "_")
        return f"user_{safe_user_id}_documents"
    
    async def ensure_collection_exists(self, user_id: str) -> str:
        """
        Ensure a collection exists for the user, create if not.
        
        Args:
            user_id: User identifier
            
        Returns:
            Collection name
        """
        collection_name = self.get_collection_name(user_id)
        print(f"[RAG] Checking if collection '{collection_name}' exists...")
        
        try:
            # Check if collection exists
            collections = self._client.get_collections()
            collection_names = [c.name for c in collections.collections]
            print(f"[RAG] Existing collections: {collection_names}")
            
            if collection_name not in collection_names:
                print(f"[RAG] Creating new collection: {collection_name}")
                # Create collection with vector configuration
                self._client.create_collection(
                    collection_name=collection_name,
                    vectors_config=VectorParams(
                        size=EMBEDDING_DIMENSION,
                        distance=Distance.COSINE
                    )
                )
                print(f"[RAG] ✅ Created Qdrant collection: {collection_name}")
            else:
                print(f"[RAG] ✅ Collection already exists: {collection_name}")
            
            return collection_name
        except Exception as e:
            print(f"[RAG] ❌ Error ensuring collection exists: {e}")
            traceback.print_exc()
            raise
    
    async def delete_collection(self, user_id: str) -> bool:
        """
        Delete a user's document collection.
        
        Args:
            user_id: User identifier
            
        Returns:
            True if deleted successfully
        """
        collection_name = self.get_collection_name(user_id)
        print(f"[RAG] Deleting collection: {collection_name}")
        
        try:
            self._client.delete_collection(collection_name)
            print(f"[RAG] ✅ Deleted Qdrant collection: {collection_name}")
            return True
        except Exception as e:
            print(f"[RAG] ❌ Error deleting collection: {e}")
            traceback.print_exc()
            return False


class DocumentProcessor:
    """
    Processes various document formats and extracts text content.
    Supports: PDF, DOCX, TXT, MD, CSV
    """
    
    SUPPORTED_EXTENSIONS = {'.pdf', '.docx', '.doc', '.txt', '.md', '.csv', '.xlsx'}
    
    @classmethod
    def is_supported(cls, filename: str) -> bool:
        """Check if file type is supported."""
        ext = os.path.splitext(filename.lower())[1]
        supported = ext in cls.SUPPORTED_EXTENSIONS
        print(f"[RAG] File '{filename}' extension '{ext}' supported: {supported}")
        return supported
    
    @classmethod
    async def extract_text(cls, file: UploadFile) -> Tuple[str, str]:
        """
        Extract text content from an uploaded file.
        
        Args:
            file: FastAPI UploadFile object
            
        Returns:
            Tuple of (extracted_text, file_type)
        """
        filename = file.filename or "unknown"
        ext = os.path.splitext(filename.lower())[1]
        
        print(f"[RAG] Extracting text from file: {filename} (extension: {ext})")
        
        # Read file content with size limit
        MAX_FILE_SIZE = 50 * 1024 * 1024  # 50 MB limit
        try:
            content = await file.read()
            file_size = len(content)
            print(f"[RAG] Read {file_size} bytes from file")
            
            if file_size > MAX_FILE_SIZE:
                raise ValueError(f"File too large: {file_size} bytes (max: {MAX_FILE_SIZE})")
            
            if file_size == 0:
                raise ValueError("File is empty")
                
            await file.seek(0)  # Reset file pointer
        except Exception as e:
            print(f"[RAG] ❌ Error reading file content: {e}")
            traceback.print_exc()
            raise
        
        try:
            if ext == '.pdf':
                print(f"[RAG] Processing as PDF...")
                text = cls._extract_from_pdf(content)
                print(f"[RAG] ✅ Extracted {len(text)} characters from PDF")
                return text, 'pdf'
            elif ext in {'.docx', '.doc'}:
                print(f"[RAG] Processing as DOCX...")
                text = cls._extract_from_docx(content)
                print(f"[RAG] ✅ Extracted {len(text)} characters from DOCX")
                return text, 'docx'
            elif ext in {'.txt', '.md'}:
                print(f"[RAG] Processing as text file...")
                text = cls._extract_from_text(content)
                print(f"[RAG] ✅ Extracted {len(text)} characters from text file")
                return text, 'text'
            elif ext == '.csv':
                print(f"[RAG] Processing as CSV...")
                text = cls._extract_from_csv(content)
                print(f"[RAG] ✅ Extracted {len(text)} characters from CSV")
                return text, 'csv'
            elif ext == '.xlsx':
                print(f"[RAG] Processing as XLSX...")
                text = cls._extract_from_xlsx(content)
                print(f"[RAG] ✅ Extracted {len(text)} characters from XLSX")
                return text, 'xlsx'
            else:
                raise ValueError(f"Unsupported file type: {ext}")
        except Exception as e:
            print(f"[RAG] ❌ Error extracting text from {filename}: {e}")
            traceback.print_exc()
            raise
    
    @staticmethod
    def _extract_from_pdf(content: bytes) -> str:
        """Extract text from PDF content."""
        print(f"[RAG] Loading PDF with pypdf... (content size: {len(content)} bytes)")
        
        try:
            from pypdf import PdfReader
        except ImportError as e:
            print(f"[RAG] ❌ Failed to import pypdf: {e}")
            raise
        
        try:
            pdf_file = io.BytesIO(content)
            reader = PdfReader(pdf_file, strict=False)  # Non-strict mode for problematic PDFs
            
            print(f"[RAG] PDF has {len(reader.pages)} pages")
            
            text_parts = []
            for i, page in enumerate(reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                        print(f"[RAG] Page {i+1}: extracted {len(page_text)} chars")
                    else:
                        print(f"[RAG] Page {i+1}: no text extracted (might be image-based)")
                except Exception as page_error:
                    print(f"[RAG] ⚠️ Page {i+1}: error extracting text: {page_error}")
                    continue
            
            if not text_parts:
                print(f"[RAG] ⚠️ No text could be extracted from PDF (might be scanned/image-based)")
                return ""
            
            return "\n\n".join(text_parts)
            
        except Exception as e:
            print(f"[RAG] ❌ Error reading PDF: {e}")
            traceback.print_exc()
            raise
    
    @staticmethod
    def _extract_from_docx(content: bytes) -> str:
        """Extract text from DOCX content."""
        print(f"[RAG] Loading DOCX with python-docx...")
        from docx import Document
        
        docx_file = io.BytesIO(content)
        doc = Document(docx_file)
        
        print(f"[RAG] DOCX has {len(doc.paragraphs)} paragraphs")
        
        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # Also extract from tables
        table_count = len(doc.tables)
        print(f"[RAG] DOCX has {table_count} tables")
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)
        
        return "\n\n".join(text_parts)
    
    @staticmethod
    def _extract_from_text(content: bytes) -> str:
        """Extract text from plain text file with encoding detection."""
        print(f"[RAG] Detecting text encoding...")
        # Detect encoding
        detected = chardet.detect(content)
        encoding = detected.get('encoding', 'utf-8') or 'utf-8'
        print(f"[RAG] Detected encoding: {encoding}")
        
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            print(f"[RAG] Fallback to utf-8 with error replacement")
            # Fallback to utf-8 with error handling
            return content.decode('utf-8', errors='replace')
    
    @staticmethod
    def _extract_from_csv(content: bytes) -> str:
        """Extract text from CSV content."""
        import csv
        
        print(f"[RAG] Processing CSV...")
        # Detect encoding
        detected = chardet.detect(content)
        encoding = detected.get('encoding', 'utf-8') or 'utf-8'
        
        text_content = content.decode(encoding, errors='replace')
        csv_file = io.StringIO(text_content)
        
        reader = csv.reader(csv_file)
        rows = []
        for row in reader:
            rows.append(" | ".join(cell.strip() for cell in row if cell.strip()))
        
        print(f"[RAG] Processed {len(rows)} CSV rows")
        return "\n".join(rows)
    
    @staticmethod
    def _extract_from_xlsx(content: bytes) -> str:
        """Extract text from Excel file."""
        print(f"[RAG] Processing XLSX...")
        from openpyxl import load_workbook
        
        xlsx_file = io.BytesIO(content)
        workbook = load_workbook(xlsx_file, read_only=True, data_only=True)
        
        print(f"[RAG] XLSX has {len(workbook.sheetnames)} sheets: {workbook.sheetnames}")
        
        text_parts = []
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            text_parts.append(f"--- Sheet: {sheet_name} ---")
            
            for row in sheet.iter_rows(values_only=True):
                row_text = " | ".join(str(cell) for cell in row if cell is not None)
                if row_text.strip():
                    text_parts.append(row_text)
        
        workbook.close()
        return "\n".join(text_parts)


class ChunkingService:
    """
    Splits documents into fixed-size chunks with overlap.
    Uses character-based chunking for simplicity and consistency.
    """
    
    def __init__(self, chunk_size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP):
        """
        Initialize chunking service.
        
        Args:
            chunk_size: Maximum characters per chunk
            overlap: Number of overlapping characters between chunks
        """
        self.chunk_size = chunk_size
        self.overlap = overlap
    
    def chunk_text(
        self,
        text: str,
        filename: str,
        file_type: str
    ) -> List[DocumentChunk]:
        """
        Split text into overlapping chunks.
        
        Args:
            text: Full document text
            filename: Source filename
            file_type: Document type
            
        Returns:
            List of DocumentChunk objects
        """
        print(f"[RAG] Chunking text of length {len(text)} for file {filename}")
        
        if not text or not text.strip():
            print(f"[RAG] Warning: Empty text, no chunks generated")
            return []
        
        # Clean and normalize text
        text = text.strip()
        text = " ".join(text.split())  # Normalize whitespace
        
        text_length = len(text)
        print(f"[RAG] Normalized text length: {text_length}")
        
        # If text is smaller than chunk size, return single chunk
        if text_length <= self.chunk_size:
            print(f"[RAG] Text fits in single chunk")
            return [DocumentChunk(
                text=text,
                filename=filename,
                file_type=file_type,
                chunk_index=0,
                total_chunks=1,
                created_at=datetime.utcnow()
            )]
        
        chunks = []
        start = 0
        chunk_index = 0
        max_iterations = (text_length // 100) + 10  # Safety limit
        iteration = 0
        
        while start < text_length and iteration < max_iterations:
            iteration += 1
            
            # Calculate end position
            end = start + self.chunk_size
            
            if end >= text_length:
                # Last chunk - take everything remaining
                chunk_text = text[start:]
                print(f"[RAG] Chunk {chunk_index + 1}: last chunk, {len(chunk_text)} chars (start={start})")
            else:
                # Try to break at a sentence boundary or word boundary
                chunk_text = text[start:end]
                
                # Look for sentence boundary near the end
                last_period = chunk_text.rfind('. ')
                last_newline = chunk_text.rfind('\n')
                last_space = chunk_text.rfind(' ')
                
                # Prefer sentence boundaries, then newlines, then spaces
                break_point = max(last_period, last_newline)
                if break_point > self.chunk_size * 0.7:  # At least 70% of chunk
                    chunk_text = chunk_text[:break_point + 1]
                elif last_space > self.chunk_size * 0.7:
                    chunk_text = chunk_text[:last_space]
                
                print(f"[RAG] Chunk {chunk_index + 1}: {len(chunk_text)} chars (start={start})")
            
            if chunk_text.strip():
                chunks.append(DocumentChunk(
                    text=chunk_text.strip(),
                    filename=filename,
                    file_type=file_type,
                    chunk_index=chunk_index,
                    total_chunks=0,  # Will be updated
                    created_at=datetime.utcnow()
                ))
                chunk_index += 1
            
            # Calculate next start position
            chunk_len = len(chunk_text)
            
            # If this is the last chunk or remaining text is small, we're done
            if start + chunk_len >= text_length:
                break
            
            # Move to next chunk with overlap, but ensure we always advance
            advancement = max(chunk_len - self.overlap, 100)  # Always advance at least 100 chars
            start += advancement
            
            # Ensure start doesn't go backwards
            if start <= chunks[-1].chunk_index if chunks else 0:
                start = (chunks[-1].chunk_index + 1) * 100 if chunks else 100
        
        if iteration >= max_iterations:
            print(f"[RAG] ⚠️ Hit max iterations limit ({max_iterations}), stopping chunking")
        
        # Update total chunks count
        total = len(chunks)
        for chunk in chunks:
            chunk.total_chunks = total
        
        print(f"[RAG] ✅ Generated {total} chunks")
        return chunks


class RAGService:
    """
    Main RAG service orchestrating document storage and retrieval.
    """
    
    def __init__(self):
        """Initialize RAG service with required components."""
        print("[RAG] Initializing RAG Service...")
        try:
            self.qdrant_manager = QdrantClientManager()
            self.document_processor = DocumentProcessor()
            self.chunking_service = ChunkingService()
            self.embedding_service = get_embedding_service()
            print("[RAG] ✅ RAG Service initialized successfully")
        except Exception as e:
            print(f"[RAG] ❌ Failed to initialize RAG Service: {e}")
            traceback.print_exc()
            raise
    
    async def store_documents(
        self,
        user_id: str,
        files: List[UploadFile]
    ) -> Dict[str, Any]:
        """
        Process and store documents in Qdrant for a user.
        
        Args:
            user_id: User identifier
            files: List of uploaded files
            
        Returns:
            Summary of stored documents
        """
        print(f"\n[RAG] ========== STORE DOCUMENTS START ==========")
        print(f"[RAG] User: {user_id}")
        print(f"[RAG] Number of files: {len(files) if files else 0}")
        
        if not files:
            print("[RAG] No files to process")
            return {"success": True, "message": "No files to process", "stored": 0}
        
        # Log file info
        for i, f in enumerate(files):
            print(f"[RAG] File {i+1}: {f.filename} (content_type: {f.content_type})")
        
        # Ensure collection exists
        print(f"\n[RAG] Step 1: Ensure collection exists...")
        try:
            collection_name = await self.qdrant_manager.ensure_collection_exists(user_id)
            print(f"[RAG] ✅ Collection ready: {collection_name}")
        except Exception as e:
            print(f"[RAG] ❌ Failed to ensure collection: {e}")
            traceback.print_exc()
            return {
                "success": False,
                "message": f"Failed to create collection: {str(e)}",
                "errors": [str(e)]
            }
        
        all_chunks: List[DocumentChunk] = []
        processed_files = []
        errors = []
        
        # Process each file
        print(f"\n[RAG] Step 2: Process files...")
        for file in files:
            filename = file.filename or "unknown"
            print(f"\n[RAG] Processing file: {filename}")
            
            # Check if file type is supported
            if not self.document_processor.is_supported(filename):
                error_msg = f"Unsupported file type: {filename}"
                print(f"[RAG] ❌ {error_msg}")
                errors.append(error_msg)
                continue
            
            try:
                # Extract text from file
                print(f"[RAG] Extracting text from {filename}...")
                text, file_type = await self.document_processor.extract_text(file)
                
                if not text or not text.strip():
                    error_msg = f"No text extracted from: {filename}"
                    print(f"[RAG] ❌ {error_msg}")
                    errors.append(error_msg)
                    continue
                
                print(f"[RAG] ✅ Extracted {len(text)} characters from {filename}")
                
                # Chunk the document
                print(f"[RAG] Chunking document...")
                chunks = self.chunking_service.chunk_text(text, filename, file_type)
                all_chunks.extend(chunks)
                processed_files.append({
                    "filename": filename,
                    "file_type": file_type,
                    "chunks": len(chunks)
                })
                
                print(f"[RAG] ✅ Processed {filename}: {len(chunks)} chunks")
                
            except Exception as e:
                error_msg = f"Error processing {filename}: {str(e)}"
                print(f"[RAG] ❌ {error_msg}")
                traceback.print_exc()
                errors.append(error_msg)
        
        if not all_chunks:
            print(f"[RAG] ❌ No chunks generated from any files")
            return {
                "success": False,
                "message": "No chunks generated from files",
                "errors": errors
            }
        
        print(f"\n[RAG] Step 3: Generate embeddings for {len(all_chunks)} chunks...")
        
        # Generate embeddings for all chunks
        try:
            chunk_texts = [chunk.text for chunk in all_chunks]
            print(f"[RAG] Calling embedding service...")
            embeddings = await self.embedding_service.generate_embeddings_batch(chunk_texts)
            
            if len(embeddings) != len(all_chunks):
                raise ValueError(f"Embedding count mismatch: {len(embeddings)} vs {len(all_chunks)}")
            
            print(f"[RAG] ✅ Generated {len(embeddings)} embeddings")
            
        except Exception as e:
            print(f"[RAG] ❌ Error generating embeddings: {e}")
            traceback.print_exc()
            return {
                "success": False,
                "message": f"Failed to generate embeddings: {str(e)}",
                "errors": errors
            }
        
        # Store in Qdrant
        print(f"\n[RAG] Step 4: Store {len(all_chunks)} chunks in Qdrant...")
        try:
            points = []
            for i, (chunk, embedding) in enumerate(zip(all_chunks, embeddings)):
                point_id = str(uuid.uuid4())
                points.append(PointStruct(
                    id=point_id,
                    vector=embedding,
                    payload={
                        "text": chunk.text,
                        "filename": chunk.filename,
                        "file_type": chunk.file_type,
                        "chunk_index": chunk.chunk_index,
                        "total_chunks": chunk.total_chunks,
                        "created_at": chunk.created_at.isoformat(),
                        "user_id": user_id
                    }
                ))
            
            print(f"[RAG] Created {len(points)} point structures")
            
            # Upsert points in batches
            batch_size = 100
            for i in range(0, len(points), batch_size):
                batch = points[i:i + batch_size]
                print(f"[RAG] Upserting batch {i//batch_size + 1} ({len(batch)} points)...")
                self.qdrant_manager.client.upsert(
                    collection_name=collection_name,
                    points=batch
                )
            
            print(f"[RAG] ✅ Stored {len(points)} chunks for user {user_id}")
            print(f"[RAG] ========== STORE DOCUMENTS END ==========\n")
            
            return {
                "success": True,
                "message": f"Successfully stored {len(points)} chunks",
                "stored": len(points),
                "files_processed": processed_files,
                "errors": errors if errors else None
            }
            
        except Exception as e:
            print(f"[RAG] ❌ Error storing in Qdrant: {e}")
            traceback.print_exc()
            return {
                "success": False,
                "message": f"Failed to store in Qdrant: {str(e)}",
                "errors": errors
            }
    
    async def retrieve_relevant_chunks(
        self,
        user_id: str,
        query: str,
        top_k: int = 5,
        score_threshold: float = 0.3
    ) -> List[RetrievedChunk]:
        """
        Retrieve relevant document chunks for a query.
        
        Args:
            user_id: User identifier
            query: Search query
            top_k: Number of chunks to retrieve
            score_threshold: Minimum similarity score
            
        Returns:
            List of retrieved chunks with scores
        """
        print(f"\n[RAG] ========== RETRIEVE CHUNKS START ==========")
        print(f"[RAG] User: {user_id}")
        print(f"[RAG] Query: {query[:100]}...")
        
        collection_name = self.qdrant_manager.get_collection_name(user_id)
        print(f"[RAG] Collection name: {collection_name}")
        
        try:
            # Check if collection exists
            print(f"[RAG] Step 1: Check if collection exists...")
            collections = self.qdrant_manager.client.get_collections()
            collection_names = [c.name for c in collections.collections]
            
            if collection_name not in collection_names:
                print(f"[RAG] No document collection found for user {user_id}")
                print(f"[RAG] ========== RETRIEVE CHUNKS END ==========\n")
                return []
            
            print(f"[RAG] ✅ Collection exists")
            
            # Generate query embedding
            print(f"[RAG] Step 2: Generate query embedding...")
            query_embedding = await self.embedding_service.generate_embedding(query)
            print(f"[RAG] ✅ Generated embedding (dim: {len(query_embedding)})")
            
            # Search in Qdrant using query_points (v1.16+ API)
            print(f"[RAG] Step 3: Query Qdrant (top_k={top_k}, threshold={score_threshold})...")
            results = self.qdrant_manager.client.query_points(
                collection_name=collection_name,
                query=query_embedding,
                limit=top_k,
                score_threshold=score_threshold
            )
            
            print(f"[RAG] ✅ Query returned {len(results.points)} results")
            
            retrieved_chunks = []
            for i, result in enumerate(results.points):
                payload = result.payload or {}
                chunk = RetrievedChunk(
                    text=payload.get("text", ""),
                    filename=payload.get("filename", "unknown"),
                    score=result.score,
                    chunk_index=payload.get("chunk_index", 0)
                )
                retrieved_chunks.append(chunk)
                print(f"[RAG] Result {i+1}: score={result.score:.4f}, file={chunk.filename}")
            
            print(f"[RAG] ========== RETRIEVE CHUNKS END ==========\n")
            return retrieved_chunks
            
        except Exception as e:
            print(f"[RAG] ❌ Error retrieving chunks: {e}")
            traceback.print_exc()
            print(f"[RAG] ========== RETRIEVE CHUNKS END ==========\n")
            return []
    
    async def delete_user_documents(self, user_id: str) -> bool:
        """
        Delete all documents for a user.
        
        Args:
            user_id: User identifier
            
        Returns:
            True if successful
        """
        print(f"[RAG] Deleting documents for user: {user_id}")
        return await self.qdrant_manager.delete_collection(user_id)
    
    def format_context_for_prompt(
        self,
        chunks: List[RetrievedChunk],
        max_length: int = 4000
    ) -> str:
        """
        Format retrieved chunks as context for LLM prompt.
        
        Args:
            chunks: List of retrieved chunks
            max_length: Maximum context length in characters
            
        Returns:
            Formatted context string
        """
        print(f"[RAG] Formatting {len(chunks)} chunks for prompt (max_length={max_length})")
        
        if not chunks:
            print(f"[RAG] No chunks to format")
            return ""
        
        context_parts = []
        total_length = 0
        
        for chunk in chunks:
            chunk_text = f"[From: {chunk.filename}]\n{chunk.text}"
            
            if total_length + len(chunk_text) > max_length:
                # Truncate if needed
                remaining = max_length - total_length
                if remaining > 100:
                    chunk_text = chunk_text[:remaining] + "..."
                    context_parts.append(chunk_text)
                break
            
            context_parts.append(chunk_text)
            total_length += len(chunk_text) + 2  # +2 for separators
        
        if not context_parts:
            return ""
        
        context = "\n\n---\n\n".join(context_parts)
        final_context = f"\n\nRelevant context from your documents:\n\n{context}"
        print(f"[RAG] ✅ Formatted context: {len(final_context)} characters")
        return final_context


# Global RAG service instance
_rag_service: Optional[RAGService] = None


def get_rag_service() -> RAGService:
    """
    Get or create global RAG service instance.
    
    Returns:
        RAGService instance
    """
    global _rag_service
    if _rag_service is None:
        print("[RAG] Creating new RAG service instance...")
        _rag_service = RAGService()
    return _rag_service
